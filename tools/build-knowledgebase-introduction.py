from __future__ import annotations

from pathlib import Path
from urllib.parse import quote

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_NAME = "MediaHedge_Knowledgebase_Introduction.docx"
VAULT_NAME = "MH Wiki"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
SLATE = "4F6272"
MUTED = "6D7782"
GOLD = "B38A2E"
LIGHT_BLUE = "E8EEF5"
PALE_GOLD = "FFF8E8"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "1F2328"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, *, name="Calibri", size=None, color=BLACK, bold=None, italic=None):
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, *, size, color=BLACK, bold=False, italic=False, before=0, after=6, line=1.25):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    style.font.italic = italic
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_paragraph(paragraph, fill: str, left_border: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if left_border:
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_border)
        borders.append(left)


def add_hyperlink(paragraph, text: str, url: str, *, color=BLUE, bold=False, italic=False):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.append(size)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def obsidian_uri(file_path: str) -> str:
    return f"obsidian://open?vault={quote(VAULT_NAME)}&file={quote(file_path, safe='')}"


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, cached, end])
    set_run_font(run, size=9, color=MUTED)


def setup_numbering(document: Document, *, kind: str, marker: str, left_dxa: int, hanging_dxa: int) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), kind)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left_dxa))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_dxa))
    ind.set(qn("w:hanging"), str(hanging_dxa))
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D3D9E1", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(document, text: str, level: int):
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(document, text: str, *, bold_lead: str | None = None):
    paragraph = document.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_link_line(document, label: str, path: str, description: str):
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(5)
    add_hyperlink(p, label, obsidian_uri(path), bold=True)
    r = p.add_run(f" - {description}")
    set_run_font(r)
    return p


def configure_document() -> tuple[Document, int, int]:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    set_style(styles["Normal"], size=11, after=6, line=1.25)
    set_style(styles["Title"], size=30, color=NAVY, bold=True, before=0, after=8, line=1.0)
    set_style(styles["Subtitle"], size=15, color=DARK_BLUE, before=0, after=6, line=1.15)
    set_style(styles["Heading 1"], size=16, color=BLUE, bold=True, before=18, after=10, line=1.0)
    set_style(styles["Heading 2"], size=13, color=BLUE, bold=True, before=14, after=7, line=1.0)
    set_style(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True, before=10, after=5, line=1.0)

    bullet_num = setup_numbering(document, kind="bullet", marker="•", left_dxa=540, hanging_dxa=270)
    decimal_num = setup_numbering(document, kind="decimal", marker="%1.", left_dxa=540, hanging_dxa=270)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("MEDIAHEDGE KNOWLEDGEBASE")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = p.add_run("\tA GUIDED TOUR")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("MediaHedge | Film-finance knowledge, connected")
    set_run_font(left, size=8.5, color=MUTED)
    right = p.add_run("\t")
    set_run_font(right, size=8.5, color=MUTED)
    add_field(p, "PAGE")

    document.core_properties.title = "Welcome to the MediaHedge Knowledgebase"
    document.core_properties.subject = "A guided introduction to MediaHedge and its film-finance wiki"
    document.core_properties.author = "MediaHedge Wiki"
    document.core_properties.keywords = "MediaHedge, film finance, knowledgebase, wiki, underwriting"
    document.core_properties.comments = "Derived orientation artifact based on the MediaHedge wiki; not independent source authority."
    return document, bullet_num, decimal_num


def build_document(output_path: Path):
    document, bullet_num, decimal_num = configure_document()

    # Editorial cover (named overrides: cover title 30 pt navy; kicker 11 pt gold).
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(98)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("WELCOME TO MEDIAHEDGE")
    set_run_font(run, size=11, color=GOLD, bold=True)

    title = document.add_paragraph("The MediaHedge Knowledgebase", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("A friendly guided tour of how film-finance credit is built, protected, serviced, and measured", style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tagline = document.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_before = Pt(18)
    tagline.paragraph_format.space_after = Pt(72)
    run = tagline.add_run("Start with the story. Follow the cash. Understand the controls.")
    set_run_font(run, size=11, color=GOLD, bold=True, italic=True)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_after = Pt(5)
    run = date.add_run("August 2026")
    set_run_font(run, size=11, color=NAVY, bold=True)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Open the blue links in this guide to jump directly into the MH Wiki vault in Obsidian.")
    set_run_font(run, size=9.5, color=MUTED, italic=True)

    document.add_page_break()

    add_heading(document, "Hello, and welcome", 1)
    add_body(
        document,
        "MediaHedge sits at the meeting point of creative production and disciplined credit. In this knowledgebase, the company is represented as a specialist film- and television-finance originator, underwriter, and servicer that helps turn complex production assets and contracts into an understandable, governable lending system.",
    )
    add_body(
        document,
        "You do not need to arrive as a film-finance expert. The wiki is designed to let you begin with the big picture, follow a topic that interests you, and keep drilling until you reach the underlying source brief.",
    )

    callout = document.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(10)
    shade_paragraph(callout, PALE_GOLD, GOLD)
    lead = callout.add_run("The big idea: ")
    set_run_font(lead, size=11.5, color=NAVY, bold=True)
    rest = callout.add_run("verify the value, size the exposure, control the cash, monitor the change, and preserve the options.")
    set_run_font(rest, size=11.5, color=NAVY)

    add_heading(document, "What you can expect to learn", 1)
    learning = [
        ("Why completion comes first", "wiki/concepts/full-financing", "See how full financing protects the path to delivery - and why a small funding gap can weaken several repayment sources at once."),
        ("How a loan earns its size", "wiki/concepts/loan-sizing", "Follow collateral eligibility, advance rates, concentration, leverage, budget, term, and liquidity constraints to the binding ceiling."),
        ("How protection becomes operational", "wiki/concepts/protection-stack", "Connect the protection stack to security, insurance, completion coverage, account control, and enforceable remedies."),
        ("How each dollar finds the lender", "wiki/concepts/cash-control-and-waterfalls", "Understand payment directions, CAMA administration, deposit-account control, waterfalls, and reconciliation without confusing their roles."),
        ("How risk evolves after closing", "wiki/concepts/monitoring-and-servicing", "Learn how budgets, collateral, contracts, collections, and covenants change state - and how monitoring turns variance into action."),
        ("How performance is judged", "wiki/concepts/financier-return-economics", "Separate stated coupon from dated cash realization, XIRR, cash multiple, expected loss, and the operating friction between gross and net return."),
    ]
    for label, path, description in learning:
        p = document.add_paragraph(style="Normal")
        apply_numbering(p, decimal_num)
        add_hyperlink(p, label, obsidian_uri(path), bold=True)
        run = p.add_run(f". {description}")
        set_run_font(run)

    add_heading(document, "The system, not a single silver bullet", 2)
    add_body(
        document,
        "The wiki repeatedly distinguishes tools that sound similar but do different jobs. A CAMA is not Article 9 account control. Production insurance is not a completion guaranty. A high coupon does not make weak collateral eligible. A large number of productions does not create diversification when the same obligor, jurisdiction, guarantor, or delivery event sits underneath them all.",
    )
    add_link_line(document, "See the Repayment and Risk Map", "wiki/syntheses/repayment-and-risk-map", "a side-by-side view of the value sources, conditions, controls, and blind spots.")

    document.add_page_break()

    add_heading(document, "Choose your own adventure", 1)
    add_body(document, "There is no single correct reading order. Pick the route that matches the question you brought with you.")

    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [1800, 3900, 3660])
    set_table_borders(table)
    headers = ["If you are...", "You will learn...", "Start here"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=10, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])

    routes = [
        ("New to film finance", "The shared language, the major repayment paths, and why delivery connects so many risks.", [("Overview", "wiki/overview"), ("Glossary", "wiki/glossary")]),
        ("A financier or credit partner", "How eligibility, sizing, concentration, governance, and dated cash returns fit together.", [("Policy rails", "wiki/syntheses/policy-rails-and-control-matrix"), ("Credit lifecycle", "wiki/syntheses/credit-lifecycle")]),
        ("Legal, closing, or operations", "How rights are created, payments are captured, exceptions are escalated, and servicing stays portable.", [("Security package", "wiki/concepts/security-package"), ("Forward-flow governance", "wiki/concepts/forward-flow-governance")]),
        ("Risk, workout, or recovery focused", "How early diagnosis preserves completion, claims, collateral, cash, and remedy choices.", [("Workouts and recoveries", "wiki/concepts/defaults-workouts-and-recoveries"), ("Protection stack", "wiki/concepts/protection-stack")]),
    ]
    for reader, learns, links in routes:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[0].paragraphs[0]
        run = p.add_run(reader)
        set_run_font(run, size=10, color=NAVY, bold=True)
        p = cells[1].paragraphs[0]
        run = p.add_run(learns)
        set_run_font(run, size=10)
        p = cells[2].paragraphs[0]
        for link_idx, (label, path) in enumerate(links):
            if link_idx:
                r = p.add_run(" | ")
                set_run_font(r, size=10, color=MUTED)
            add_hyperlink(p, label, obsidian_uri(path), bold=True)
    set_table_geometry(table, [1800, 3900, 3660])

    add_heading(document, "A quick tour of the protection stack", 1)
    stack = [
        ("Eligibility and verification", "Only documented, enforceable, independently supported value enters the borrowing base."),
        ("Sizing and structural cushion", "Advance rates, leverage ceilings, concentration limits, reserves, and tenor define exposure before pricing."),
        ("Priority and cash control", "Assignments, liens, notices, controlled accounts, waterfalls, and reconciliation make rights usable."),
        ("Completion and risk transfer", "Full financing, completion protection, and insurance address specific failure modes - never every failure mode."),
        ("Surveillance and enforcement", "Monitoring, covenants, triggers, decision rights, and recovery playbooks keep options alive after closing."),
    ]
    for label, detail in stack:
        p = document.add_paragraph(style="Normal")
        apply_numbering(p, bullet_num)
        r = p.add_run(f"{label}. ")
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(detail)
        set_run_font(r)
    add_link_line(document, "Explore the full Protection Stack", "wiki/concepts/protection-stack", "the hub that connects underwriting, closing, servicing, and recovery controls.")

    document.add_page_break()

    add_heading(document, "How the knowledgebase works", 1)
    add_body(
        document,
        "The vault is intentionally simple. Raw Word briefs preserve the evidence. Source-summary pages explain what each brief says and where its limits are. Concept pages maintain durable explanations. Synthesis pages connect several concepts into lifecycle, policy, and risk views. The index keeps the whole system navigable.",
    )

    layers = [
        ("Raw sources", "Immutable snapshots and hashes", "raw/manifest"),
        ("Compiled wiki", "Source pages, concepts, entities, syntheses, and the glossary", "index"),
        ("Operating schema", "Rules for evidence, ingestion, contradictions, links, and maintenance", "AGENTS"),
    ]
    for title_text, explanation, path in layers:
        p = document.add_paragraph(style="Normal")
        apply_numbering(p, decimal_num)
        add_hyperlink(p, title_text, obsidian_uri(path), bold=True)
        r = p.add_run(f" - {explanation}.")
        set_run_font(r)

    add_heading(document, "A knowledgebase that shows its homework", 2)
    evidence = [
        ("Source-backed", "A claim stated directly in one or more cited briefs."),
        ("Synthesis", "A conclusion formed by connecting several source-backed claims."),
        ("Inference", "Reasoning that goes beyond direct wording and is labeled as such."),
        ("Open question", "A gap that still needs evidence, authority, or a decision."),
    ]
    for label, detail in evidence:
        p = document.add_paragraph(style="Normal")
        apply_numbering(p, bullet_num)
        r = p.add_run(f"{label}: ")
        set_run_font(r, bold=True, color=DARK_BLUE)
        r = p.add_run(detail)
        set_run_font(r)

    caution = document.add_paragraph()
    caution.paragraph_format.left_indent = Inches(0.18)
    caution.paragraph_format.right_indent = Inches(0.12)
    caution.paragraph_format.space_before = Pt(8)
    caution.paragraph_format.space_after = Pt(10)
    shade_paragraph(caution, LIGHT_GRAY, BLUE)
    r = caution.add_run("A useful guardrail: ")
    set_run_font(r, color=NAVY, bold=True)
    r = caution.add_run("the wiki preserves the source corpus's internal policy statements, but undated limits remain marked for review. Executed transaction documents and current legal, tax, insurance, and program authority still control.")
    set_run_font(r, color=NAVY)

    add_heading(document, "Your five-minute starting route", 1)
    quick_start = [
        ("Open the welcoming Wiki Index", "index", "Meet the knowledgebase and see the complete map."),
        ("Read the Executive Overview", "wiki/overview", "Absorb the lending model in one connected story."),
        ("Walk the Credit Lifecycle", "wiki/syntheses/credit-lifecycle", "Follow a loan from screening through recovery."),
        ("Keep the Glossary nearby", "wiki/glossary", "Translate the specialized language as you go."),
        ("Check the Research Backlog", "wiki/operations/research-backlog", "See which questions still need policy, legal, operational, or performance evidence."),
    ]
    for label, path, description in quick_start:
        add_link_line(document, label, path, description)

    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(18)
    closing.paragraph_format.space_after = Pt(8)
    r = closing.add_run("Welcome in. Pick a thread, follow the links, and let the system reveal itself one control at a time.")
    set_run_font(r, size=12.5, color=GOLD, bold=True, italic=True)

    provenance = document.add_paragraph()
    provenance.alignment = WD_ALIGN_PARAGRAPH.CENTER
    provenance.paragraph_format.space_before = Pt(10)
    provenance.paragraph_format.space_after = Pt(0)
    r = provenance.add_run("Derived orientation artifact based on the MediaHedge wiki's twelve-source corpus. It introduces existing knowledge and adds no independent evidence.")
    set_run_font(r, size=8.5, color=MUTED, italic=True)

    document.save(output_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    target = root / OUTPUT_NAME
    build_document(target)
    print(target)
